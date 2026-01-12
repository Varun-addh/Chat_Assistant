from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Header, Response
from typing import Optional, List
from fastapi.responses import StreamingResponse
from datetime import datetime

from app.schemas import (
	CreateSessionResponse, 
	QuestionIn, 
	AnswerOut, 
	SessionHistory, 
	QnA, 
	SessionList, 
	SessionSummary, 
	UpdateSessionTitleRequest
)
from app.services.session_manager import get_session_manager, SessionManager, SessionState
from app.services.llm_service import get_llm_service
from app.services.architecture_generator import get_architecture_generator, ArchitectureViewType
from app.services.code_evaluation_service import evaluate_code
from app.utils.security import verify_api_key
from app.utils.audit import auditor
from app.config import settings
from app.middleware.auth import get_user_id_from_request
from fastapi import Request
import asyncio
import logging
import uuid

from app.utils.story_contract import extract_mermaid_first, enforce_story_contract, sanitize_mermaid_subset
from app.utils.mermaid_sanitizer import MermaidSanitizer

from app.database import get_db_context
from app.utils.usage_tracking import track_api_usage, track_session
from app.utils.event_logging import track_event

logger = logging.getLogger(__name__)

# Use intelligent provider selection with "copilot" feature flag for AI Assistant/Chat
llm_service = get_llm_service(feature="copilot")

router = APIRouter()


def _validate_specificity(text: str, system_description: str) -> tuple[bool, list[str]]:
	"""Check if explanation is specific enough for FAANG-level interviews.
	
	Returns (is_specific, list_of_issues).
	"""
	import re
	
	issues = []
	text_lower = text.lower()
	
	# Check for generic platitudes (indicates shallow thinking)
	generic_phrases = [
		"business logic", "process the request", "store the data",
		"handle the workflow", "manage the system", "provides scalability",
		"ensures reliability", "improves performance", "handle requests"
	]
	
	generic_count = sum(1 for phrase in generic_phrases if phrase in text_lower)
	if generic_count > 3:
		issues.append(f"Too many generic phrases ({generic_count}): use actual tech names and mechanisms")
	
	# Specificity should be domain-agnostic and avoid buzzword forcing.
	# We look for:
	# - concrete architecture nouns / components
	# - at least one metric (or explicit comparison)
	# - at least one reliability mechanism
	tech_names = [
		r'\b(redis|memcached|postgres|postgresql|mysql|cassandra|mongodb|dynamodb|bigtable|elasticsearch|opensearch)\b',
		r'\b(kafka|pubsub|pub/sub|rabbitmq|sqs|kinesis)\b',
		r'\b(s3|gcs|object storage|cdn|cloudfront|cloudflare)\b',
	]
	arch_nouns = [
		r'\b(api gateway|gateway|load balancer|waf|cdn|edge)\b',
		r'\b(cache|caching|ttl|eviction|lru)\b',
		r'\b(database|db|index|indexing|search)\b',
		r'\b(queue|broker|worker|job|scheduler|dlq)\b',
		r'\b(replication|failover|partitioning|shard|sharding|consistency)\b',
	]
	metric_patterns = [
		r'\b(p50|p95|p99)\b',
		r'\b\d+\s*(ms|milliseconds?|s|sec|seconds?|m|min|minutes?|hrs?|hours?)\b',
		r'\b\d+(?:\.\d+)?\s*(rps|qps|tps|req/s|req/sec)\b',
		r'\b\d+(?:\.\d+)?\s*(kb|mb|gb|tb|pb)\b',
		r'\b\d+(?:\.\d+)?%\b',
	]
	reliability_patterns = [
		r'\b(retry|backoff|timeout|circuit breaker|fallback|idempotent|idempotency)\b',
		r'\b(rate limit|throttle|shed load)\b',
		r'\b(reconcile|reconciliation|replay|dead letter|dlq)\b',
	]

	has_tech = any(re.search(p, text, re.IGNORECASE) for p in tech_names)
	has_arch = any(re.search(p, text, re.IGNORECASE) for p in arch_nouns)
	has_metric = any(re.search(p, text, re.IGNORECASE) for p in metric_patterns)
	has_comparison = bool(re.search(r'\b(vs|versus|compared to|instead of|chose\s+.*\s+over)\b', text_lower))
	has_reliability = any(re.search(p, text, re.IGNORECASE) for p in reliability_patterns)

	if not (has_tech or has_arch):
		issues.append("Missing concrete components (e.g., cache/DB/queue/CDN/gateway) or recognizable technologies")

	if not (has_metric or has_comparison):
		issues.append("Missing quantified reasoning (metrics with units or an explicit trade-off comparison)")

	if not has_reliability:
		issues.append("Missing reliability mechanisms (retry/backoff/timeout/DLQ/fallback/idempotency)")
	
	return len(issues) == 0, issues


def _sanitize_mermaid_from_llm(mermaid_code: str) -> str:
	"""Remove CSS/style/HTML artifacts that LLM might inject into Mermaid code.

	Delegates to the single-source Mermaid sanitization pipeline.
	"""
	return MermaidSanitizer.sanitize_from_llm(mermaid_code)


async def _auto_evaluate_if_code(session_manager: SessionManager, session_id: str, question: str, answer: str, api_key: Optional[str] = None) -> None:
	"""Auto-evaluate if the answer contains code blocks."""
	import re
	
	# Look for code blocks in the answer
	code_pattern = r'```(\w+)?\n(.*?)\n```'
	matches = re.findall(code_pattern, answer, re.DOTALL)
	
	if not matches:
		return
	
	# Get the most substantial code block
	best_code = ""
	best_lang = "python"  # Default to Python for coding questions
	for lang, code in matches:
		if len(code.strip()) > len(best_code.strip()):
			best_code = code.strip()
			best_lang = lang or "python"  # Default to Python if no language specified
	
	if not best_code:
		return
	
	# Auto-trigger evaluation in background
	try:
		# Get conversation context
		session_state = await session_manager.get_required(session_id)
		conversation_context = ""
		if session_state.qna:
			recent_qna = session_state.qna[-2:] if len(session_state.qna) >= 2 else session_state.qna
			for item in recent_qna:
				conversation_context += f"Q: {item.get('question', '')}\nA: {item.get('answer', '')}\n"
		
		# Run evaluation
		await evaluate_code(question, best_code, best_lang, conversation_context, api_key=api_key)
		
		# Log auto-evaluation
		await auditor.log({
			"type": "auto_evaluation",
			"session_id": session_id,
			"question": question,
			"language": best_lang,
			"auto_triggered": True,
		})
	except Exception as e:
		# Don't fail the main request if evaluation fails
		await auditor.log({
			"type": "auto_evaluation_error",
			"session_id": session_id,
			"error": str(e),
		})


@router.post("/session", response_model=CreateSessionResponse)
async def create_session(request: Request):
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	state = await manager.create_session()

	# Track sessions only for authenticated users (guests remain file-based only)
	user = getattr(request.state, "user", None)
	if user:
		with get_db_context() as db:
			track_session(db, user, state.session_id, session_type="qa")
	return CreateSessionResponse(session_id=state.session_id)


@router.post("/question")
async def submit_question(
	payload: QuestionIn,
	request: Request,
	response: Response,
	x_api_key: Optional[str] = Header(None, alias="X-API-Key"),
	x_gemini_key: Optional[str] = Header(None, alias="X-Gemini-Key"),
	authorization: Optional[str] = Header(None, alias="Authorization")
):
	# 1. Extract and clean headers (handle empty strings/undefined/null from frontend)
	def _clean(v: Optional[str]) -> Optional[str]:
		t = (v or "").strip()
		if not t:
			return None
		if t.lower() in {"null", "undefined", "none"}:
			return None
		return t

	groq_key = _clean(x_api_key)
	gemini_key = _clean(x_gemini_key)

	# IMPORTANT:
	# - If the request is authenticated, Authorization is assumed to be JWT (not an LLM key).
	# - If unauthenticated, accept Authorization only when it looks like a real LLM key.
	user = getattr(request.state, "user", None)
	if user is None and not groq_key and authorization:
		auth_val = _clean(authorization)
		if auth_val:
			if auth_val.lower().startswith("bearer "):
				token = auth_val.split(" ", 1)[1].strip() if " " in auth_val else ""
			else:
				token = auth_val
			# Only treat Authorization as an LLM key if it matches Groq/Gemini key shapes
			if token.startswith("gsk_") or token.startswith("AIza") or token.startswith("ATza"):
				groq_key = token

	# 2. Decide demo vs registered for this endpoint.
	# Product requirement: demo mode should use Groq (cost-capped), not Gemini.
	# We treat the request as demo only when:
	# - not authenticated, AND
	# - no bridge key headers were supplied.
	is_demo = (user is None) and (not groq_key) and (not gemini_key)

	# 3. Select key to use.
	# If the user provided a Gemini key in Bridge Settings, we prefer it (higher quality).
	# Otherwise, use Groq key if provided.
	api_key = gemini_key or groq_key
	
	# 4. Fallback to server/env keys ONLY if user provided NO keys in Bridge Settings
	if not api_key:
		if is_demo:
			# DEMO PATH: always prefer Groq
			if settings.is_demo_key_pool_enabled():
				try:
					from app.services.demo_key_pool import demo_key_pool
					demo_key = demo_key_pool.get_key()
					if demo_key and demo_key.startswith("gsk_"):
						api_key = demo_key
						logger.info("🔧 [DEMO] Using Groq demo key pool")
				except Exception:
					# Fall back to env Groq key below
					pass

			if not api_key and settings.groq_api_key:
				api_key = settings.groq_api_key
				logger.info("🔧 [DEMO] Using environment GROQ_API_KEY")
		else:
			# REGISTERED PATH: use effective provider selection
			effective = settings.get_effective_provider(feature="copilot")
			if effective == "gemini" and settings.gemini_api_key:
				api_key = settings.gemini_api_key
				logger.info("🔧 Using environment GEMINI_API_KEY")
			elif effective == "groq" and settings.groq_api_key:
				api_key = settings.groq_api_key
				logger.info("🔧 Using environment GROQ_API_KEY")
			else:
				api_key = settings.gemini_api_key or settings.groq_api_key
				if api_key:
					logger.info(f"🔧 Using fallback environment key for {effective}")

	# 5. Final check: if still no key, raise error
	if not api_key:
		if is_demo:
			raise HTTPException(
				status_code=503,
				detail="Demo is temporarily unavailable (Groq is not configured).",
			)
		raise HTTPException(
			status_code=401,
			detail="No active API key. Please add your Groq or Gemini API key in Bridge Settings to continue.",
		)
	
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	recovered_from_session_id: Optional[str] = None
	reused_existing_session: bool = False

	def _is_valid_session_id(value: str) -> bool:
		v = (value or "").strip()
		if not v:
			return False
		if v.lower() in {"undefined", "null", "none"}:
			return False
		try:
			uuid.UUID(v)
			return True
		except Exception:
			return False

	async def _get_most_recent_session() -> Optional[SessionState]:
		"""Best-effort: return the most recently updated session for this user."""
		try:
			items = await manager.list_sessions()
		except Exception:
			return None
		for it in items:
			sid = it.get("session_id")
			if not sid:
				continue
			state = await manager.get(sid)
			if state is not None:
				return state
		return None
	
	# Guard: if the frontend sent a bogus session_id, reuse latest session instead of
	# creating a new one for every message (prevents sidebar spam).
	if not _is_valid_session_id(payload.session_id):
		recovered_from_session_id = payload.session_id
		recent = await _get_most_recent_session()
		if recent is not None:
			state = recent
			payload.session_id = state.session_id
			reused_existing_session = True
			logger.info(f"Reused most recent session {state.session_id} due to invalid session_id='{recovered_from_session_id}'")
		else:
			state = await manager.create_session()
			payload.session_id = state.session_id
			logger.info(f"Created new session because session_id was invalid and no recent session exists: {state.session_id}")
	else:
		try:
			state = await manager.get(payload.session_id)
			if state is None:
				# Auto-recovery: prefer reusing the most recent session if it was updated recently.
				recovered_from_session_id = payload.session_id
				recent = await _get_most_recent_session()
				if recent is not None:
					state = recent
					payload.session_id = state.session_id
					reused_existing_session = True
					logger.info(
						f"Session {recovered_from_session_id} not found. Reusing most recent session {state.session_id}"
					)
				else:
					logger.info(f"Session {recovered_from_session_id} not found. Creating new session for recovery.")
					state = await manager.create_session()
					payload.session_id = state.session_id
					logger.info(f"Auto-created new session: {state.session_id} (recovered from {recovered_from_session_id})")
		except KeyError:
			# Fallback: create new session
			recovered_from_session_id = payload.session_id
			state = await manager.create_session()
			payload.session_id = state.session_id
			logger.info(f"Created new session due to KeyError: {state.session_id}")

	# Expose the effective session_id to the client (helps UI stay stable across refresh/tabs)
	response.headers["X-Stratax-Session-Id"] = payload.session_id
	if recovered_from_session_id and recovered_from_session_id != payload.session_id:
		response.headers["X-Stratax-Session-Recovered"] = "1"
		response.headers["X-Stratax-Old-Session-Id"] = recovered_from_session_id
		if reused_existing_session:
			response.headers["X-Stratax-Session-Reused"] = "1"

	if not payload.question.strip():
		raise HTTPException(status_code=400, detail="Empty question")

	# --- Telemetry spine: request received (safe, low-cost, no raw text by default) ---
	if getattr(settings, "enable_event_logging", True):
		try:
			with get_db_context() as db:
				track_event(
					db,
					user_id=user_id,
					session_id=payload.session_id,
					event_type="chat_prompt_received",
					question_text=payload.question,
					extra={
						"stream": bool(payload.stream),
						"saved_to_history": bool(payload.save_to_history),
						"architecture_mode": payload.architecture_mode,
						"style_mode": getattr(payload, "style_mode", None),
						"tone": getattr(payload, "tone", None),
						"layout": getattr(payload, "layout", None),
					},
				)
		except Exception:
			# Never fail the request due to analytics.
			pass

	# Auto-detect architecture mode choice if user replies with "single" or "multi"
	# This handles the case where user is responding to the architecture choice prompt
	if not payload.architecture_mode:
		q_lower = payload.question.strip().lower()
		# Check if this is a simple mode selection response
		if q_lower in ["single", "single-view", "single view", "1"]:
			payload.architecture_mode = "single"
			logger.info(f"🎯 Auto-detected architecture mode choice: single")
			# Retrieve the original question from conversation history
			if state.qna and len(state.qna) > 0:
				last_qna = state.qna[-1]
				if isinstance(last_qna, dict) and "question" in last_qna:
					original_question = last_qna["question"]
					logger.info(f"📝 Retrieved original question from history: {original_question[:100]}")
					payload.question = original_question
				elif isinstance(last_qna, QnA):
					original_question = last_qna.question
					logger.info(f"📝 Retrieved original question from history: {original_question[:100]}")
					payload.question = original_question
		elif q_lower in ["multi", "multi-view", "multi view", "multiview", "2"]:
			payload.architecture_mode = "multi-view"
			logger.info(f"🎯 Auto-detected architecture mode choice: multi-view")
			# Retrieve the original question from conversation history
			if state.qna and len(state.qna) > 0:
				last_qna = state.qna[-1]
				if isinstance(last_qna, dict) and "question" in last_qna:
					original_question = last_qna["question"]
					logger.info(f"📝 Retrieved original question from history: {original_question[:100]}")
					payload.question = original_question
				elif isinstance(last_qna, QnA):
					original_question = last_qna.question
					logger.info(f"📝 Retrieved original question from history: {original_question[:100]}")
					payload.question = original_question

	# Deterministic identity/developer attribution answers (never call an LLM)
	# This is a belt-and-suspenders guard in addition to the LLMService short-circuit.
	if llm_service._is_identity_question(payload.question):
		identity_answer = llm_service._identity_response_text(payload.question)
		if getattr(settings, "enable_event_logging", True):
			try:
				with get_db_context() as db:
					track_event(
						db,
						user_id=user_id,
						session_id=payload.session_id,
						event_type="chat_identity_guard",
						question_text=payload.question,
						extra={"stream": bool(payload.stream)},
					)
			except Exception:
				pass
		# Proof-of-handling: inspect this in DevTools -> Network -> Response Headers
		response.headers["X-Stratax-Guard"] = "identity"
		response.headers["X-Stratax-App"] = getattr(settings, "app_name", "Stratax AI")
		if payload.stream:
			async def _identity_event_gen():
				# SSE: prefix each line with 'data: '
				safe = identity_answer.replace("\n", "\ndata: ")
				yield f"data: {safe}\n\n"
				# Persist to history if requested
				if payload.save_to_history:
					await manager.append_qna(payload.session_id, payload.question, identity_answer)
				await auditor.log({
					"type": "qna",
					"session_id": payload.session_id,
					"question": payload.question,
					"answer": identity_answer,
					"saved_to_history": payload.save_to_history,
				})
				yield "event: end\n\n"
			# StreamingResponse doesn't automatically inherit headers from `response`
			return StreamingResponse(
				_identity_event_gen(),
				media_type="text/event-stream",
				headers={
					"X-Stratax-Guard": "identity",
					"X-Stratax-App": getattr(settings, "app_name", "Stratax AI"),
					"X-Stratax-Session-Id": payload.session_id,
					**(
						{"X-Stratax-Session-Recovered": "1", "X-Stratax-Old-Session-Id": recovered_from_session_id}
						if recovered_from_session_id and recovered_from_session_id != payload.session_id
						else {}
					),
				},
			)

		if payload.save_to_history:
			await manager.append_qna(payload.session_id, payload.question, identity_answer)
		await auditor.log({
			"type": "qna",
			"session_id": payload.session_id,
			"question": payload.question,
			"answer": identity_answer,
			"saved_to_history": payload.save_to_history,
		})
		return AnswerOut(answer=identity_answer, created_at=datetime.utcnow(), truncated=False)

	# Read any stored profile for this session
	profile_text = state.profile_text

	# --- ARCHITECTURE DETECTION ---
	# Smart pattern-based detection (no LLM call needed - saves time and cost)
	q_lower = payload.question.lower()
	
	# Pattern 1: Explicit system design keywords
	has_explicit = any(kw in q_lower for kw in settings.architecture_detection_explicit_keywords)
	
	# Pattern 2: "Design/Build X" pattern (catches "design twitter", "build a cache", etc.)
	has_design_verb = any(kw in q_lower for kw in settings.architecture_detection_design_verbs)
	
	# Pattern 3: System components/concepts (indicates architecture discussion)
	has_system_concepts = (
		sum(
			[
				any(kw in q_lower for kw in settings.architecture_detection_system_concepts_scale),
				any(kw in q_lower for kw in settings.architecture_detection_system_concepts_data),
				any(kw in q_lower for kw in settings.architecture_detection_system_concepts_infra),
			]
		)
		>= 2
	)  # Need at least 2 system concepts
	
	# Pattern 4: Exclude coding problems
	is_code_problem = any(kw in q_lower for kw in settings.architecture_detection_code_problem_keywords)
	
	# Trigger if: explicit keywords OR design verb OR multiple system concepts, AND not a code problem
	is_architecture = (has_explicit or has_design_verb or has_system_concepts) and not is_code_problem

	# If system design detected AND user hasn't chosen a mode yet, ask for preference.
	arch_mode = payload.architecture_mode
	if is_architecture and not arch_mode:
		# IMPORTANT: Do NOT leak a magic trigger string into chat text.
		# Instead, return a structured UI hint so the frontend can render a selector.
		# The client should re-submit /api/question with architecture_mode set.
		logger.info("🏗️ System design detected with no architecture_mode: returning ui_action chooser")
		if getattr(settings, "enable_event_logging", True):
			try:
				with get_db_context() as db:
					track_event(
						db,
						user_id=user_id,
						session_id=payload.session_id,
						event_type="chat_architecture_mode_prompted",
						question_text=payload.question,
						extra={"detected": True},
					)
			except Exception:
				pass
		return AnswerOut(
			answer="",
			created_at=datetime.utcnow(),
			truncated=False,
			ui_action="choose_architecture_mode",
			ui_payload={
				"default": "multi-view",
				"options": [
					{"id": "single", "label": "Single View", "description": "One comprehensive architecture diagram"},
					{"id": "multi-view", "label": "Multi View", "description": "5 focused diagrams: overview, flow, data, deployment, observability"},
				],
			},
		)

	# Route based on user's choice
	if is_architecture and arch_mode == "multi-view":
		logger.info(f"🏗️ Detected system design question: {payload.question}")
		
		# View titles and descriptions for each architecture layer
		view_info = {
			"SYSTEM_OVERVIEW": ("🏗️ System Overview", "High-level components and how they connect"),
			"REQUEST_FLOW": ("🔄 Request Flow", "Step-by-step journey of a user request through the system"),
			"DATA_MODEL": ("🗄️ Data & Storage", "Databases, caches, and how data is persisted"),
			"DEPLOYMENT": ("📦 Deployment Architecture", "Infrastructure, scaling, and cloud services"),
			"OBSERVABILITY": ("📊 Observability & Monitoring", "Metrics, logs, alerts, and system health"),
		}
		
		# Define the generator for architecture responses
		async def architecture_stream_gen():
			try:
				# 1. Initialize services
				arch_gen = get_architecture_generator()

				def _sanitize_view_explanation(text: str) -> str:
					"""Make LLM output deterministic for the UI.

					We already emit an H2 title per view. This strips any accidental
					H1/H2 headings and removes stray 'Bottom line:' lines that often
					show up inside sections.
					"""
					import re
					if not text:
						return ""
					out_lines: list[str] = []
					for raw_line in text.splitlines():
						line = raw_line.rstrip()
						# Remove top-level headings (caller provides them)
						if re.match(r"^\s*#{1,2}\s+", line):
							continue
						# Remove stray 'Bottom line' inside layers (keep final summaries clean)
						if re.match(r"^\s*([\-*]\s*)?(\*\*?)?bottom line(\*\*?)?:", line, flags=re.IGNORECASE):
							continue
						out_lines.append(raw_line)
					return "\n".join(out_lines).strip()

				def _safe_ascii(text: str) -> str:
					# Keep UI stable and avoid non-ASCII issues.
					return (text or "").encode("ascii", errors="ignore").decode("ascii", errors="ignore")

				# 2. Select views: keep multi-view deterministic (the product promise is 5 views).
				views_to_gen = [
					ArchitectureViewType.SYSTEM_OVERVIEW,
					ArchitectureViewType.REQUEST_FLOW,
					ArchitectureViewType.DATA_MODEL,
					ArchitectureViewType.DEPLOYMENT,
					ArchitectureViewType.OBSERVABILITY,
				]
				
				def _fallback_view(view_name: str, view_type: ArchitectureViewType, err: Exception | None = None) -> tuple[str, str]:
					"""Return (mermaid_code, explanation) placeholders for a view.

					We use this when the LLM provider rate-limits or errors, so the
					frontend still receives a complete 5-view package.
					"""
					err_msg = str(err).replace("\n", " ") if err else "Unknown error"
					mermaid = (
						"flowchart TD\n"
						"  Client[Client] --> API[API]\n"
						"  API --> Compute[Compute]\n"
						"  Compute --> Store[(Storage)]\n"
					)
					# Prefer the same 5-layer scaffold the prompts enforce.
					# (Internal method, but keeps formatting consistent.)
					layer_scaffold = ""
					try:
						layer_scaffold = arch_gen._build_layer_explanations(view_type)  # type: ignore[attr-defined]
					except Exception:
						layer_scaffold = "### Layer 1 - Client & Auth\n- What happens: (generation unavailable)\n\n### Final Summary\nGeneration unavailable."

					if view_name == "SYSTEM_OVERVIEW":
						explanation = (
							"- Generation degraded due to an LLM/provider error (showing safe fallback).\n"
							"- This view still shows the core building blocks and their primary links.\n"
							"- Retry later if you want richer domain-specific details.\n"
							"- Error: " + _safe_ascii(err_msg) + "\n"
							"- The remaining views may also degrade if the provider is rate-limited.\n"
							"Goal: Provide a stable multi-view package even under partial failures."
						)
					else:
						explanation = (
							"NOTE: LLM generation failed for this view; showing the canonical 5-layer template.\n"
							"Error: " + _safe_ascii(err_msg) + "\n\n" + layer_scaffold
						)

					return mermaid, explanation

				# 3. Generate each view with diagram + story-driven explanation
				full_response = f"# System Design: {payload.question}\n\n"
				
				for idx, view_type in enumerate(views_to_gen):
					is_last_view = (idx == len(views_to_gen) - 1)
					
					# Get view title and description
					view_name = view_type.name if hasattr(view_type, 'name') else str(view_type)
					title, desc = view_info.get(view_name, (f"📐 {view_name}", "Architecture view"))
					
					# Send title
					yield f"data: ## {title}\n\n"
					
					prompt_data = arch_gen.get_view_prompt(view_type, payload.question)
					# prompt_data['system_prompt'] already contains a strict output contract.
					base_prompt = f"{prompt_data['user_prompt']}"
					combined_prompt = base_prompt

					# Generate view with up to 1 retry; if provider errors, fall back per-view.
					response_text = ""
					mermaid_code = ""
					explanation = ""
					last_issues: List[str] = []

					try:
						# Keep retries tightly bounded to prevent quota storms.
						for attempt in range(2):
							response_text, _ = await llm_service.generate_answer(
								question=combined_prompt,
								system_prompt=prompt_data['system_prompt'],
								api_key=api_key,
								apply_auto_overrides=False,
								allow_provider_fallback=(not is_demo),
							)

							# Parse response: Mermaid diagram first, then explanation
							mermaid_code, raw_explanation = extract_mermaid_first(response_text or "")
							if not mermaid_code:
								mermaid_code = "flowchart TD\n  Error[No diagram generated]"
							explanation = raw_explanation
							explanation = _sanitize_view_explanation(explanation)
							explanation = enforce_story_contract(view_name, payload.question, explanation)
							explanation = _safe_ascii(explanation)
							mermaid_code = sanitize_mermaid_subset(mermaid_code)

							# Validate specificity (skip for SYSTEM_OVERVIEW which is intentionally brief)
							if view_name != "SYSTEM_OVERVIEW":
								is_specific, specificity_issues = _validate_specificity(explanation, payload.question)
								if not is_specific and attempt < 1:
									logger.warning(f"[{view_name}] Content not specific enough (attempt {attempt+1}): {specificity_issues}")
									combined_prompt = (
									f"MAKE IT MORE SPECIFIC. Previous answer was too generic: {', '.join(specificity_issues)}.\n\n"
									"Hard rules:\n"
									"- Stay within ONE system. Do not introduce unrelated systems.\n"
									"- Keep the cloud stack consistent (pick one provider; do not mix AWS/GCP/Azure unless explicitly asked).\n"
									"- If DOMAIN HINTS are present in the prompt, you MUST follow them.\n"
									"- Avoid buzzword dumping. Use advanced tech ONLY if it is truly required and you justify it in 'Why this layer exists'.\n\n"
									"You MUST include (specific to THIS system):\n"
									"- Concrete components with meaningful names (e.g., CDN, origin, catalog service, playback service, cache, DB).\n"
									"- 2-3 real metrics with units (p95 latency ms, throughput rps, storage TB/day, cache TTL, availability %).\n"
									"- 1-2 domain-appropriate mechanisms (examples: idempotency keys + retries, optimistic locking, cache invalidation, ABR via HLS/DASH, DRM/license, fanout notifications).\n"
									"- At least ONE realistic failure mode per layer with mitigation (retry/backoff, DLQ, circuit breaker, fallback, reconciliation).\n\n"
									"BAD: 'caching layer stores data'\n"
									"GOOD: 'Redis: 3 shards, LRU eviction, 10m TTL for playback manifests; cache warmed on trending titles; stale-while-revalidate for spikes'\n\n"
									+ base_prompt
								)
									continue  # Retry with specificity guidance

							# Validate complexity; if too complex, retry with a simplification instruction
							validation = arch_gen.validate_diagram_complexity(
								mermaid_code,
								view_type,
								max_nodes=prompt_data.get("max_nodes"),
								max_edges=prompt_data.get("max_edges"),
							)
							if validation.get("valid"):
								break

							last_issues = validation.get("issues") or []
							logger.warning(f"Complexity warning for {view_type}: {last_issues}")
							if attempt < 1:
								combined_prompt = (
								f"SIMPLIFY AND REGENERATE. Previous diagram violated constraints: {last_issues}.\n"
								f"You MUST reduce to <= {prompt_data.get('max_nodes')} nodes and <= {prompt_data.get('max_edges')} edges.\n"
								"Keep ONLY the critical path for this view. Remove optional components.\n"
								"Do NOT use subgraphs unless absolutely necessary.\n\n"
								"Also: Explanation must follow the OUTPUT CONTRACT exactly (no Key Highlights, no tables, no deep dives).\n\n"
								+ base_prompt
							)

					except Exception as e:
						logger.warning(f"[{view_name}] View generation failed, using fallback: {e}")
						mermaid_code, explanation = _fallback_view(view_name, view_type, e)

					# If still invalid after retries, proceed (renderer will show placeholder if needed)
					
					# Validate it starts with a valid Mermaid diagram type
					first_line = (mermaid_code.split('\n')[0].strip().lower() if mermaid_code else "")
					valid_starts = ['flowchart', 'graph', 'sequencediagram', 'classdiagram', 'statediagram']
					if not any(first_line.startswith(vs) for vs in valid_starts):
						logger.error(f"Invalid Mermaid diagram - doesn't start with valid type. First line: {first_line}")
						mermaid_code = "flowchart TD\n  Error[Invalid diagram generated]"
					
					# CRITICAL: Strip any CSS/style content that LLM might have added (breaks rendering)
					mermaid_code = _sanitize_mermaid_from_llm(mermaid_code)
					
					# Log the mermaid code for debugging
					logger.info(f"Generated Mermaid code ({len(mermaid_code)} chars): {mermaid_code[:200]}...")
						
					# Validate (log only)
					validation = arch_gen.validate_diagram_complexity(mermaid_code, view_type)
					if not validation.get("valid"):
						logger.warning(f"Complexity warning for {view_type}: {validation.get('issues')}")

					# Send diagram (SSE-safe: prefix every line with 'data: ')
					diagram_block = f"```mermaid\n{mermaid_code}\n```"
					safe_diagram = diagram_block.replace("\n", "\ndata: ")
					yield f"data: {safe_diagram}\n\n"
					
					# Stream explanation (includes code/tips if last view)
					safe_explanation = explanation.replace(chr(10), chr(10) + 'data: ')
					yield f"data: {safe_explanation}\n\n"
					
					# Append to full response for history
					full_response += f"## {title}\n\n```mermaid\n{mermaid_code}\n```\n\n{explanation}\n\n"
				
				# Save to history (code + tips already included in last view)
				if payload.save_to_history:
					await manager.append_qna(state.session_id, payload.question, full_response)
					
				yield "event: end\n\n"
				
			except Exception as e:
				logger.error(f"Architecture generation failed: {e}")
				err_msg = str(e).replace(chr(10), " ")
				yield f"data: ⚠️ **Generation failed:** {err_msg}\n\n"
				yield "event: end\n\n"

		if payload.stream:
			return StreamingResponse(architecture_stream_gen(), media_type="text/event-stream")
		else:
			# Non-streaming fallback
			chunks = []
			async for chunk in architecture_stream_gen():
				if chunk.startswith("data: "):
					chunks.append(chunk[6:].replace("\ndata: ", "\n"))
			final_ans = "".join(chunks)
			if payload.save_to_history:
				await manager.append_qna(state.session_id, payload.question, final_ans)
			return AnswerOut(answer=final_ans, created_at=datetime.utcnow(), truncated=False)
	
	# Handle single comprehensive architecture (legacy single-diagram behavior)
	elif is_architecture and arch_mode == "single":
		logger.info(f"🏗️ Generating SINGLE comprehensive architecture for: {payload.question}")
		
		# IMPORTANT: "single" here means the legacy single-diagram system-design behavior
		# powered by LLMService's built-in system-design overrides (the "old architecture thing").
		# We do NOT run the newer multi-view architecture generator or SINGLE story-contract clamping.
		logger.info("📊 Using legacy SINGLE architecture path (LLMService auto system-design overrides)")

		if payload.stream:
			async def single_architecture_stream_gen():
				collected: list[str] = []
				previous_qna = state.qna
				# Quick feedback + consistent header for the UI
				yield "data: ## Single-View Architecture\n\n"
				yield "data: Generating architecture...\n\n"

				async for chunk in llm_service.stream_answer(
					payload.question,
					payload.system_prompt,
					profile_text=profile_text,
					previous_qna=previous_qna[-10:] if previous_qna else None,
					style_mode=payload.style_mode,
					tone=payload.tone,
					layout=payload.layout,
					variability=payload.variability,
					seed=payload.seed,
					api_key=api_key,
					# Legacy behavior: allow LLMService to apply its auto system-design overrides.
					apply_auto_overrides=True,
					allow_provider_fallback=(not is_demo),
				):
					collected.append(chunk)
					safe_chunk = chunk.replace('\n', '\ndata: ')
					yield f"data: {safe_chunk}\n\n"

				if payload.save_to_history:
					await manager.append_qna(
						state.session_id,
						payload.question,
						"## Single-View Architecture\n\n" + "".join(collected),
					)
				yield "event: end\n\n"

			return StreamingResponse(single_architecture_stream_gen(), media_type="text/event-stream")
		else:
			previous_qna = state.qna
			answer, truncated = await llm_service.generate_answer(
				payload.question,
				payload.system_prompt,
				profile_text=profile_text,
				previous_qna=previous_qna[-10:] if previous_qna else None,
				style_mode=payload.style_mode,
				tone=payload.tone,
				layout=payload.layout,
				variability=payload.variability,
				seed=payload.seed,
				api_key=api_key,
				# Legacy behavior: allow LLMService to apply its auto system-design overrides.
				apply_auto_overrides=True,
				allow_provider_fallback=(not is_demo),
			)
			final_ans = ("## Single-View Architecture\n\n" + (answer or "").strip()).strip()
			if payload.save_to_history:
				await manager.append_qna(state.session_id, payload.question, final_ans)
			return AnswerOut(answer=final_ans, created_at=datetime.utcnow(), truncated=truncated)
		

	
	# Regular non-architecture questions
	if payload.stream:
		async def event_gen():
			collected: list[str] = []
			# Provide recent QnA as context for follow-ups
			previous_qna = state.qna
			async for chunk in llm_service.stream_answer(
				payload.question,
				payload.system_prompt,
				profile_text=profile_text,
				previous_qna=previous_qna,
				style_mode=payload.style_mode,
				tone=payload.tone,
				layout=payload.layout,
				variability=payload.variability,
				seed=payload.seed,
				api_key=api_key,
				allow_provider_fallback=(not is_demo),
			):
				collected.append(chunk)
				# Fix: Proper SSE encoding for multi-line chunks to prevent text loss/corruption
				# EventSource spec requires 'data: ' prefix for every line of the data payload
				safe_chunk = chunk.replace('\n', '\ndata: ')
				yield f"data: {safe_chunk}\n\n"
			# On stream end, persist the full answer
			full_answer = "".join(collected)
			if payload.save_to_history:
				await manager.append_qna(state.session_id, payload.question, full_answer)
			
			await auditor.log({
				"type": "qna",
				"session_id": state.session_id,
				"question": payload.question,
				"answer": full_answer,
				"saved_to_history": payload.save_to_history,
			})

			if getattr(settings, "enable_event_logging", True):
				try:
					with get_db_context() as db:
						track_event(
							db,
							user_id=user_id,
							session_id=state.session_id,
							event_type="chat_answer_generated",
							question_text=payload.question,
							extra={
								"stream": True,
								"saved_to_history": bool(payload.save_to_history),
								"answer_len": len(full_answer or ""),
							},
						)
				except Exception:
					pass
			
			# Auto-evaluate if response contains code
			asyncio.create_task(_auto_evaluate_if_code(manager, state.session_id, payload.question, full_answer, api_key))
			
			yield "event: end\n\n"
		return StreamingResponse(event_gen(), media_type="text/event-stream")

	# Provide recent QnA as context for follow-ups
	previous_qna = state.qna
	try:
		answer, truncated = await llm_service.generate_answer(
			payload.question,
			payload.system_prompt,
			profile_text=profile_text,
			previous_qna=previous_qna,
			style_mode=payload.style_mode,
			tone=payload.tone,
			layout=payload.layout,
			variability=payload.variability,
			seed=payload.seed,
			api_key=api_key,
			allow_provider_fallback=(not is_demo),
		)
	except Exception as e:
		logger.error("❌ LLM generate_answer failed: %s", str(e)[:300])
		raise HTTPException(
			status_code=502,
			detail=(
				"LLM provider failed. Please verify your GEMINI_API_KEY / GROQ_API_KEY (or add a key in Bridge Settings) and try again."
			),
		)
	
	if payload.save_to_history:
		await manager.append_qna(state.session_id, payload.question, answer)

	# Usage tracking (Phase 2 billing/analytics). For now tokens/cost are not computed here.
	with get_db_context() as db:
		track_api_usage(
			db,
			getattr(request.state, "user", None),
			feature="copilot",
			endpoint=str(request.url.path),
			metadata={"stream": False, "saved_to_history": bool(payload.save_to_history)},
			guest_user_id=user_id if user_id.startswith("guest_") else None,
		)
		
	await auditor.log({
		"type": "qna",
		"session_id": state.session_id,
		"question": payload.question,
		"answer": answer,
		"saved_to_history": payload.save_to_history,
	})

	if getattr(settings, "enable_event_logging", True):
		try:
			with get_db_context() as db:
				track_event(
					db,
					user_id=user_id,
					session_id=state.session_id,
					event_type="chat_answer_generated",
					question_text=payload.question,
					extra={
						"stream": False,
						"saved_to_history": bool(payload.save_to_history),
						"truncated": bool(truncated),
						"answer_len": len(answer or ""),
					},
				)
		except Exception:
			pass
	
	# Auto-evaluate if response contains code
	asyncio.create_task(_auto_evaluate_if_code(manager, state.session_id, payload.question, answer, api_key))
	
	return AnswerOut(answer=answer, created_at=datetime.utcnow(), truncated=truncated)


@router.post("/upload_profile")
async def upload_profile(
    file: UploadFile = File(...),
    session_id: str = Form(...),
    request: Request = None,
):
	# Get user-specific manager
	user_id = get_user_id_from_request(request) if request else "guest_unknown"
	manager = get_session_manager(user_id)
	
	# Ensure session exists
	try:
		await manager.get_required(session_id)
	except KeyError:
		raise HTTPException(status_code=404, detail="Session not found. Create one via POST /api/session and reuse its session_id.")

	# Determine how to read the file; support text and pdf minimally
	filename = (file.filename or "").lower()
	content_type = (file.content_type or "").lower()

	# Read bytes
	data = await file.read()

	text: str = ""
	try:
		if filename.endswith(".txt") or content_type.startswith("text/"):
			text = data.decode("utf-8", errors="ignore")
		elif filename.endswith(".md"):
			text = data.decode("utf-8", errors="ignore")
		elif filename.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
			# Handle Word documents (.docx)
			try:
				from docx import Document
				import io
				doc = Document(io.BytesIO(data))
				# Extract all paragraphs
				text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
				# Also extract text from tables
				for table in doc.tables:
					for row in table.rows:
						for cell in row.cells:
							if cell.text.strip():
								text += "\n" + cell.text
			except ImportError:
				raise HTTPException(status_code=415, detail="Word document support not available. Please install python-docx or upload a PDF/text file.")
			except Exception as e:
				raise HTTPException(status_code=415, detail=f"Unable to read Word document: {str(e)}. Please try converting to PDF or text format.")
		elif filename.endswith(".pdf") or content_type == "application/pdf":
			# Lazy import to avoid heavy dependency at startup
			try:
				from app.utils.text_extract import extract_text_from_pdf  # type: ignore
			except Exception:
				raise HTTPException(status_code=415, detail="PDF support not available. Please install pdfminer.six or upload a text/markdown file.")
			text = extract_text_from_pdf(data)
		else:
			# Fallback: try utf-8
			text = data.decode("utf-8", errors="ignore")
	except UnicodeDecodeError:
		raise HTTPException(status_code=415, detail="Unable to decode file. Please upload a UTF-8 text, markdown, PDF, or Word (.docx) file.")

	if not text.strip():
		raise HTTPException(status_code=400, detail="Uploaded file appears empty.")

	await manager.set_profile_text(session_id, text)
	await auditor.log({
		"type": "profile_upload",
		"session_id": session_id,
		"filename": file.filename,
		"bytes": len(data),
	})

	return {"status": "ok", "characters": len(text)}


@router.get("/sessions", response_model=SessionList)
async def list_sessions(request: Request):
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	items_raw = await manager.list_sessions()
	items = [
		SessionSummary(
			session_id=i["session_id"],
			title=i["title"],
			last_update=datetime.fromisoformat(i["last_update"]),
			qna_count=i["qna_count"],
		)
		for i in items_raw
	]
	return SessionList(items=items)


@router.delete("/session/{session_id}")
async def delete_session(session_id: str, request: Request):
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	deleted = await manager.delete_session(session_id)
	if not deleted:
		raise HTTPException(status_code=404, detail="Session not found")
	return {"status": "ok", "deleted": True}


@router.delete("/history/{session_id}")
async def clear_history(session_id: str, request: Request):
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	try:
		await manager.clear_history(session_id)
	except KeyError:
		raise HTTPException(status_code=404, detail="Session not found")
	return {"status": "ok"}


@router.delete("/history/{session_id}/{index}")
async def delete_qna_item(session_id: str, index: int, request: Request):
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	try:
		await manager.remove_qna(session_id, index)
	except KeyError:
		raise HTTPException(status_code=404, detail="Session not found")
	except IndexError:
		raise HTTPException(status_code=400, detail="QnA index out of range")
	return {"status": "ok"}


@router.put("/session/{session_id}/title")
async def update_session_title(session_id: str, payload: UpdateSessionTitleRequest, request: Request):
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	try:
		await manager.set_session_title(session_id, payload.title)
	except KeyError:
		raise HTTPException(status_code=404, detail="Session not found")
	return {"status": "ok", "title": payload.title}


@router.get("/session/{session_id}/chat", response_model=SessionHistory)
async def get_session_chat(session_id: str, request: Request, response: Response):
	"""Get chat history for a session (different from Search Intelligence history)"""
	user_id = get_user_id_from_request(request) or "guest_unknown"
	manager = get_session_manager(user_id)
	try:
		state = await manager.get_required(session_id)
	except KeyError:
		# Backward compatibility: earlier builds stored guest sessions under "guest_unknown".
		# After introducing stable guest IDs (guest_<hash>), old tabs may still hold a legacy
		# session_id. For guests only, try loading from legacy bucket and migrate forward.
		if user_id.startswith("guest_") and user_id != "guest_unknown":
			legacy = get_session_manager("guest_unknown")
			legacy_state = await legacy.get(session_id)
			if legacy_state is not None:
				# Migrate session to current guest bucket so future loads succeed.
				manager._sessions[session_id] = legacy_state  # type: ignore[attr-defined]
				manager._save(legacy_state, force=True)  # type: ignore[attr-defined]
				await legacy.delete_session(session_id)
				response.headers["X-Stratax-Session-Legacy-Migrated"] = "1"
				state = legacy_state
				# Continue to build response below
			else:
				raise HTTPException(status_code=404, detail="Session not found")
		else:
			raise HTTPException(status_code=404, detail="Session not found")
	
	items = [
		QnA(
			question=i["question"],
			answer=i["answer"],
			created_at=datetime.fromisoformat(i["created_at"]),
		)
		for i in state.qna
	]
	return SessionHistory(session_id=session_id, items=items)
